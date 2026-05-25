"""Reusable style helpers for naskah_revisi.docx.

COMPACT typography applied here to fit Pilmapres limits
(Body ≤ 15 pp, Lampiran ≤ 5 pp). See Page Budget Constraints in the plan.
"""
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_NAME = "Times New Roman"
BODY_SIZE = Pt(11)         # compact (was 12)
HEADING1_SIZE = Pt(13)     # compact (was 14)
HEADING2_SIZE = Pt(12)
CAPTION_SIZE = Pt(10)      # compact (was 11)
TABLE_CELL_SIZE = Pt(9)    # compact (was 11)
LINE_SPACING = 1.15        # compact (was 1.5)
TABLE_HEADER_FILL = "1F4E79"  # dark blue header

def set_run(run, size=BODY_SIZE, bold=False, italic=False, color=None):
    run.font.name = FONT_NAME
    run.font.size = size
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    # Force eastAsia font (Word quirk for non-Latin)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)

def add_body_paragraph(doc, text, justify=True, first_line_indent=None):
    if first_line_indent is None:
        first_line_indent = Cm(1.0)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    if first_line_indent:
        p.paragraph_format.first_line_indent = first_line_indent
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run(run)
    return p

def add_heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text.upper())
    set_run(run, size=HEADING1_SIZE, bold=True)
    return p

def add_heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    set_run(run, size=HEADING2_SIZE, bold=True, italic=True)
    return p

def add_caption(doc, text, align_center=True):
    p = doc.add_paragraph()
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    set_run(run, size=CAPTION_SIZE, bold=True)
    return p

def add_table(doc, header, rows, col_widths_cm=None):
    """Build a styled table with dark-blue header and white text."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(header))
    tbl.style = "Table Grid"
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Header row
    for i, h in enumerate(header):
        cell = tbl.rows[0].cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), TABLE_HEADER_FILL)
        tcPr.append(shd)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(h)
        set_run(run, size=TABLE_CELL_SIZE, bold=True, color="FFFFFF")
    # Data rows
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            cell = tbl.rows[r].cells[c]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(val))
            set_run(run, size=TABLE_CELL_SIZE)
    if col_widths_cm:
        for col_idx, w in enumerate(col_widths_cm):
            for row in tbl.rows:
                row.cells[col_idx].width = Cm(w)
    return tbl

def add_numbered_list(doc, items, indent_cm=1.0):
    for i, text in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(indent_cm)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = LINE_SPACING
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(f"{i}. {text}")
        set_run(run)

def add_bulleted_list(doc, items, indent_cm=1.0):
    for text in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(indent_cm)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = LINE_SPACING
        run = p.runs[0] if p.runs else p.add_run("")
        run.text = text
        set_run(run)

def add_page_break(doc):
    doc.add_page_break()

def add_image(doc, path, width_cm=14, center=True):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Cm(width_cm))
    return p
