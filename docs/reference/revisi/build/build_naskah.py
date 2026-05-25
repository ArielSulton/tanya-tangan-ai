"""Build PENSyarat AI revised naskah (.docx)."""
import sys
from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parent.parent  # docs/reference/revisi/
ASSETS = ROOT / "assets"
OUTPUT = ROOT / "output" / "naskah_revisi.docx"
TITLE = (
    "PENSyarat AI: Platform Pemahaman Kosakata Abstrak Berbasis "
    "Sistem Isyarat Bahasa Indonesia dan Kecerdasan Buatan untuk "
    "Siswa SDLB-B Tingkat Pemula"
)

sys.path.insert(0, str(Path(__file__).resolve().parent))
from styles import (  # noqa: E402
    add_body_paragraph, add_heading1, add_heading2, add_caption,
    add_table, add_numbered_list, add_bulleted_list, add_page_break,
    add_image,
)

def section_cover(doc):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from styles import add_image, set_run
    add_image(doc, str(ASSETS / "cover_pens_logo.png"), width_cm=8)
    for _ in range(2):
        doc.add_paragraph()
    for line in ["PRODUK INOVATIF",
                 "PEMILIHAN MAHASISWA BERPRESTASI 2026",
                 "PROGRAM DIPLOMA"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        set_run(r, size=Pt(14), bold=True)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE)
    set_run(r, size=Pt(12), bold=True)
    for _ in range(4):
        doc.add_paragraph()
    for line in ["OLEH:", "MOCHAMMAD ARIEL SULTON", "3323600054"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        set_run(r, bold=True)
    for _ in range(4):
        doc.add_paragraph()
    for line in ["POLITEKNIK ELEKTRONIKA NEGERI SURABAYA", "SURABAYA", "2026"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        set_run(r, bold=True)

def section_pengesahan(doc):
    from styles import set_run
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("LEMBAR PENGESAHAN")
    set_run(r, size=Pt(14), bold=True)
    # Page intentionally minimal — physical signature page

def section_pengantar(doc):
    from styles import add_body_paragraph, set_run, add_numbered_list
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("KATA PENGANTAR")
    set_run(r, size=Pt(14), bold=True)
    doc.add_paragraph()
    add_body_paragraph(doc,
        f"Puji syukur penulis haturkan kepada Tuhan Yang Maha Esa atas segala "
        f"rahmat dan karunia-Nya sehingga penulis dapat menyelesaikan Naskah "
        f"Produk Inovatif yang berjudul \"{TITLE}\". Penyusunan "
        f"naskah ini merupakan bagian dari rangkaian keikutsertaan dalam "
        f"Pemilihan Mahasiswa Berprestasi (Pilmapres) Program Diploma Tahun "
        f"2026. Naskah ini telah direvisi berdasarkan masukan dari tim "
        f"pendamping LLDIKTI, Ibu Ida, dan Bapak Adam guna memperkuat "
        f"rumusan masalah, uji efektivitas pembelajaran, kerincian dataset "
        f"gestur, aksesibilitas antarmuka, mitigasi risiko, perlindungan "
        f"data siswa, model keberlanjutan, dan tabel outcome before–after."
    )
    add_body_paragraph(doc, "Pada kesempatan ini, penulis menyampaikan terima kasih sebesar-besarnya kepada:")
    add_numbered_list(doc, [
        "Bapak Dr. Ir. Arif Irwansyah, S.T., M.Eng selaku Direktur Politeknik "
        "Elektronika Negeri Surabaya, atas bimbingan dan dukungannya yang "
        "telah memberikan kami motivasi dan arah dalam mengejar inovasi.",
        "Bapak Kholid Fathoni, S.Kom, M.T. selaku Wakil Direktur Bidang "
        "Kemahasiswaan dan Alumni Politeknik Elektronika Negeri Surabaya yang "
        "selalu memacu dan mendukung kami, untuk selalu berpartisipasi.",
        "Bapak/Ibu Dosen Pembina Mawapres Politeknik Elektronika Negeri "
        "Surabaya yang senantiasa memberikan bimbingan dan arahan dalam "
        "penyelesaian Produk Inovatif ini.",
        "Tim pendamping LLDIKTI, Ibu Ida, dan Bapak Adam atas masukan revisi "
        "yang memperkuat kualitas naskah dan produk PENSyarat AI.",
        "Serta, teman-teman yang telah bekerja sama dalam penyelesaian "
        "Produk Inovatif ini, yang tidak dapat penulis sebutkan satu persatu.",
    ])
    add_body_paragraph(doc,
        "Penulis menyadari bahwa naskah ini masih jauh dari kesempurnaan. "
        "Oleh karena itu, penulis mengharapkan kritik dan saran yang "
        "membangun demi penyempurnaan karya ini. Semoga platform PENSyarat "
        "AI dapat memberikan manfaat nyata bagi peningkatan kualitas "
        "pendidikan siswa tunarungu di Indonesia."
    )
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Surabaya, 25 Mei 2026\n\n\nPenulis")
    set_run(r)

def section_daftar_isi(doc):
    """Static daftar isi reflecting the post-revision structure.
    NOTE: page numbers below are placeholders — re-verify after final build."""
    from styles import set_run
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DAFTAR ISI")
    set_run(r, size=Pt(14), bold=True)
    doc.add_paragraph()
    entries = [
        ("LEMBAR PENGESAHAN", "i"),
        ("KATA PENGANTAR", "ii"),
        ("DAFTAR ISI", "iii"),
        ("DAFTAR GAMBAR", "iv"),
        ("DAFTAR TABEL", "iv"),
        ("1. LINGKUP PEMBAHASAN", "1"),
        ("   1.1 Lingkungan Penerima Manfaat dan Kerangka Analisis", "1"),
        ("   1.2 Pihak Terdampak sebagai Penerima Manfaat", "3"),
        ("   1.3 Pernyataan Teori Utama", "3"),
        ("2. IDENTIFIKASI POTENSI DAN KEBUTUHAN LINGKUNGAN", "3"),
        ("   2.1 Potensi Lingkungan", "3"),
        ("   2.2 Situasi dan Kebutuhan Lingkungan", "4"),
        ("   2.3 Rumusan Masalah", "5"),
        ("3. RUMUSAN TARGET PEMBANGUNAN", "5"),
        ("   3.1 Tujuan Pembangunan dan Target SMART", "5"),
        ("   3.2 Manfaat Solusi", "7"),
        ("4. ANALISIS CARA PENCAPAIAN TARGET", "7"),
        ("   4.1 Alternatif Solusi", "7"),
        ("   4.2 Pemilihan Solusi Terbaik", "8"),
        ("   4.3 Kelayakan Teknis PENSyarat AI", "8"),
        ("       4.3.1 Pengujian Akurasi Pengenalan Gestur SIBI", "8"),
        ("       4.3.2 Pengujian Fungsionalitas Sistem", "10"),
        ("       4.3.3 Rencana Uji Efektivitas Pembelajaran (Pre-test/Post-test)", "10"),
        ("       4.3.4 Pengujian Aksesibilitas Antarmuka", "11"),
        ("   4.4 Proyeksi Outcome: Perbandingan Sebelum dan Sesudah PENSyarat AI", "12"),
        ("5. PENJABARAN RENCANA KERJA", "13"),
        ("   5.1 Tahapan Utama (Pendekatan Strategis)", "13"),
        ("   5.2 Tahap dan Langkah Kegiatan", "13"),
        ("   5.3 Penjadwalan", "14"),
        ("   5.4 Mitigasi Risiko Implementasi", "14"),
        ("6. PENJABARAN INFORMASI TAMBAHAN", "15"),
        ("   6.1 Analisis Keunikan dan Orisinalitas Produk", "15"),
        ("   6.2 Struktur Organisasi Pelaksana", "16"),
        ("   6.3 Rencana Anggaran dan Sumber Pendanaan", "17"),
        ("   6.4 Proyeksi Jangka Panjang dan Mitra Pemangku Kepentingan", "17"),
        ("   6.5 Kesimpulan Dampak Inovasi", "18"),
        ("   6.6 Perlindungan Data Siswa", "18"),
        ("   6.7 Manfaat Jangka Panjang", "19"),
        ("7. VISUALISASI SOLUSI", "20"),
        ("   7.1 Desain Visualisasi Gagasan (SaHaBaT)", "20"),
        ("DAFTAR PUSTAKA", "21"),
        ("LAMPIRAN", "23"),
    ]
    for title, page in entries:
        p = doc.add_paragraph()
        p.paragraph_format.tab_stops.add_tab_stop(Cm(16), alignment=WD_TAB_ALIGNMENT.RIGHT)
        r1 = p.add_run(title)
        set_run(r1)
        r2 = p.add_run(f"\t{page}")
        set_run(r2)
    # Daftar Gambar & Tabel — short tables on a fresh page
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DAFTAR GAMBAR")
    set_run(r, size=Pt(14), bold=True)
    for txt, pg in [("Gambar 1. Alur Kerja Sistem PENSyarat AI", "9"),
                    ("Gambar 2. Visualisasi Roadmap dari PENSyarat AI", "18"),
                    ("Gambar 3. Desain Visualisasi Gagasan (SaHaBaT)", "20")]:
        p = doc.add_paragraph()
        p.paragraph_format.tab_stops.add_tab_stop(Cm(16), alignment=WD_TAB_ALIGNMENT.RIGHT)
        r = p.add_run(f"{txt}\t{pg}")
        set_run(r)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DAFTAR TABEL")
    set_run(r, size=Pt(14), bold=True)
    tabel_list = [
        ("Tabel 1. Analisis Solusi", "4"),
        ("Tabel 2. Affinity Diagram PENSyarat AI", "6"),
        ("Tabel 3. Target Pembangunan dengan Indikator SMART", "6"),
        ("Tabel 4. Hasil Pengujian Akurasi Pengenalan Gestur SIBI", "9"),
        ("Tabel 5. Hasil Pengujian Fungsionalitas Sistem", "10"),
        ("Tabel 6. Rancangan Pengujian Efektivitas Pre-test dan Post-test", "11"),
        ("Tabel 7. Penerapan Prinsip Aksesibilitas", "12"),
        ("Tabel 8. Proyeksi Outcome Sebelum vs Sesudah PENSyarat AI", "12"),
        ("Tabel 9. Identifikasi Risiko dan Strategi Mitigasi", "14"),
        ("Tabel 10. Skema Keberlanjutan dan Kemitraan", "17"),
        ("Tabel 11. Key Difference Matrix", "15"),
    ]
    for txt, pg in tabel_list:
        p = doc.add_paragraph()
        p.paragraph_format.tab_stops.add_tab_stop(Cm(16), alignment=WD_TAB_ALIGNMENT.RIGHT)
        r = p.add_run(f"{txt}\t{pg}")
        set_run(r)
def section_bab1(doc):
    from styles import add_heading1, add_heading2, add_body_paragraph
    add_heading1(doc, "1. Lingkup Pembahasan")
    add_heading2(doc, "1.1 Lingkungan Penerima Manfaat dan Kerangka Analisis")
    add_body_paragraph(doc,
        "Pendidikan merupakan hak dasar setiap manusia sebagaimana diamanatkan "
        "dalam Undang-Undang Dasar 1945 Pasal 31 ayat (1) serta Undang-Undang "
        "Nomor 20 Tahun 2003 tentang Sistem Pendidikan Nasional. Hak ini "
        "berlaku tanpa terkecuali, termasuk bagi penyandang disabilitas. "
        "Namun dalam praktiknya, kelompok penyandang tunarungu masih "
        "menghadapi hambatan signifikan dalam mengakses pendidikan yang "
        "berkualitas dan setara (WHO, 2023). Hambatan tersebut tidak "
        "semata-mata bersifat komunikatif, tetapi menyentuh aspek yang lebih "
        "mendasar: pemahaman kosakata.")
    # COMPACTED (page budget): four-statistic paragraphs merged into one
    add_body_paragraph(doc,
        "Data WHO (2023) mencatat sekitar 1,5 miliar orang di dunia "
        "mengalami gangguan pendengaran (proyeksi 2,5 miliar pada 2050); "
        "di Indonesia, Long Form SP2020 BPS (2024) mencatat 923 ribu "
        "penduduk mengalami kesulitan mendengar berat dan Survei "
        "Kesehatan Indonesia (Kemenkes, 2024) menempatkan prevalensi "
        "disabilitas pendengaran pada 0,4% populasi dengan hanya 4,1% "
        "yang menggunakan alat bantu dengar. Pada level pendidikan, "
        "Kemendikdasmen (2025) mencatat 162.806 siswa SLB pada TA "
        "2024/2025, dengan Jawa Timur menempati posisi kedua nasional "
        "(22.644 siswa).")
    add_body_paragraph(doc,
        "Permasalahan mendasar yang dihadapi siswa tunarungu, khususnya di "
        "jenjang sekolah dasar luar biasa tipe B (SDLB-B), bukan hanya "
        "keterbatasan komunikasi verbal, melainkan kesenjangan kosakata "
        "(vocabulary gap) yang lebar dan semakin melebar seiring "
        "bertambahnya usia. Penelitian Sarchet dkk. (2014) dalam Journal of "
        "Postsecondary Education and Disability menunjukkan bahwa siswa "
        "tunarungu rata-rata berada di persentil ke-19 kemampuan kosakata, "
        "jauh tertinggal dari rekan-rekan mereka yang mendengar di "
        "persentil ke-65. Convertino dkk. (2014) dalam Journal of Deaf "
        "Studies and Deaf Education menemukan bahwa kesenjangan ini melebar "
        "secara signifikan di sekitar usia 8–9 tahun, bertepatan dengan "
        "jenjang kelas 3–4 SDLB, ketika tuntutan bahasa meningkat karena "
        "hadirnya kata-kata abstrak dan konsep yang lebih kompleks.")
    # === NEW: Bu Ida feedback — abstract-vocab differentiator paragraph ===
    add_body_paragraph(doc,
        "Berbeda dari intervensi visual yang telah ada sebelumnya, PENSyarat AI difokuskan pada penjembatanan "
        "kesenjangan pada ranah kata abstrak. Tinjauan terhadap karya "
        "inovatif dan riset terdahulu pada pembelajaran kosakata berbasis "
        "visual untuk siswa tunarungu — sebagaimana dirangkum oleh Mat "
        "Yasin & Mohamad (2024) dalam SHS Web of Conferences serta Kadir "
        "(2021) dalam Aksara: Jurnal Ilmu Pendidikan Nonformal — "
        "menunjukkan bahwa intervensi visual yang sudah ada dominan "
        "menyasar kata benda konkret (contoh: \"apel\", \"kucing\", \"pohon\") "
        "melalui media foto atau flashcard. Padahal, kesenjangan pemahaman "
        "yang paling tajam pada siswa SDLB-B justru muncul pada kata "
        "abstrak dan kata keterangan derajat (contoh: \"sedikit\", \"sangat\", "
        "\"terlalu\", \"sebelum\", \"sesudah\") — area yang belum tersentuh "
        "secara sistematis. PENSyarat AI mengisi celah ini melalui "
        "representasi visual kontekstual khusus (ilustrasi mandiri yang "
        "melambangkan fungsi kata serta komparasi visual berdampingan) "
        "yang dibangun di atas pengenalan gestur SIBI real-time.")
    # === END NEW ===
    add_body_paragraph(doc,
        "Akar dari kesenjangan ini adalah keterbatasan incidental learning: "
        "anak-anak yang mendengar secara tidak sadar menyerap ribuan "
        "kosakata baru dari percakapan di sekitar mereka setiap hari, "
        "sementara anak tunarungu tidak memiliki akses terhadap sumber "
        "pembelajaran pasif ini (Convertino dkk., 2014; ASHA, 2024). "
        "Akibatnya, sebagaimana dikemukakan Haliza dkk. (2020) dalam "
        "Metabasa: Jurnal Bahasa, Sastra dan Pembelajaran, anak tunarungu "
        "cenderung hanya menguasai kata-kata benda konkret (seperti "
        "\"apel\", \"kucing\", \"buku\") dan mengalami kesulitan signifikan "
        "dalam memahami kata abstrak (seperti \"sangat\", \"terlalu\", "
        "\"sedikit\", \"sebelum\", dan \"sesudah\"). Temuan serupa "
        "dikemukakan oleh Nurizae dan Maimunah (2021) dalam Jurnal "
        "Basicedu, yang menyatakan bahwa anak tunarungu di Indonesia "
        "\"miskin dalam perbendaharaan kata sehingga kesulitan "
        "mengekspresikan bahasa dan bicaranya, serta mengalami kesulitan "
        "memahami kata-kata yang bersifat abstrak dan kata yang "
        "mengandung kiasan.\"")
    add_body_paragraph(doc,
        "Kondisi ini dikonfirmasi langsung melalui survei lapangan yang "
        "penulis lakukan di dua sekolah SLB-B di Kota Surabaya: SLB-B "
        "Karya Mulia Wonokromo (berdiri sejak 1954, salah satu sekolah "
        "tunarungu tertua di Indonesia) dan SLB Aditama Gebang Surabaya. "
        "Berdasarkan hasil wawancara dengan para guru di kedua sekolah "
        "tersebut, diperoleh pengakuan yang konsisten bahwa siswa SDLB-B "
        "seringkali tidak memahami makna dari kata-kata yang mereka "
        "gunakan, terutama kata-kata abstrak dan kata keterangan derajat "
        "dalam mata pelajaran Bahasa Indonesia. Guru menyatakan bahwa "
        "siswa sering hafal cara mengisyaratkan suatu kata dalam SIBI, "
        "namun tidak memahami makna konseptualnya, terutama karena tidak "
        "ada media visual yang memadai untuk menjelaskan perbedaan makna "
        "antar kata keterangan (misalnya, perbedaan antara \"sedikit besar\" "
        "dan \"sangat besar\"). Hasil wawancara selengkapnya dapat dilihat "
        "pada Lampiran 6.")
    add_body_paragraph(doc,
        "Di sisi lain, riset kependidikan konsisten menunjukkan bahwa "
        "pendekatan visual merupakan strategi paling efektif untuk "
        "pembelajaran kosakata siswa tunarungu, mengingat visi (penglihatan) "
        "adalah jalur penerimaan informasi utama mereka (TTAC Online, 2024; "
        "University of Minnesota CEHD, 2020). Penelitian Kadir (2021) di "
        "SLB Negeri Kota Gorontalo membuktikan bahwa intervensi media "
        "gambar meningkatkan penguasaan kosakata siswa tunarungu dari "
        "16,67% menjadi 100% hanya dalam dua siklus pembelajaran. Kajian "
        "literatur Mat Yasin dan Mohamad (2024) dalam SHS Web of "
        "Conferences juga mengonfirmasi bahwa penggunaan visual aids "
        "bersama bahasa isyarat secara konsisten lebih efektif "
        "dibandingkan pengajaran bahasa isyarat semata.")
    add_body_paragraph(doc,
        "Kemajuan kecerdasan buatan kini membuka peluang untuk menghadirkan "
        "pendekatan visual tersebut secara digital dan interaktif. "
        "Teknologi pengenalan gestur berbasis visi komputer seperti "
        "MediaPipe (Google, 2020) dan TensorFlow.js kini mampu berjalan "
        "langsung di peramban web tanpa instalasi, memungkinkan siswa "
        "SDLB-B memasukkan kata melalui gestur SIBI yang sudah mereka "
        "kuasai. Model bahasa besar (LLM) seperti LLaMA 3.3 via ChatGroq "
        "memungkinkan sistem untuk menghasilkan penjelasan kontekstual "
        "ketika kata belum tersedia dalam basis data visual.")
    add_body_paragraph(doc,
        "Berdasarkan fakta-fakta di atas, penulis mengembangkan platform "
        "digital bernama PENSyarat AI, sebuah platform pemahaman kosakata "
        "visual untuk siswa SDLB-B. Siswa menginput kata melalui gestur "
        "SIBI via kamera, dan platform menampilkan: (1) gambar objek nyata "
        "untuk kata konkret (misalnya: foto kucing untuk kata \"kucing\"); "
        "(2) representasi visual kontekstual untuk kosakata abstrak, baik "
        "berupa ilustrasi mandiri yang melambangkan fungsi kata (misalnya: "
        "visualisasi gestur menunjuk untuk mendeskripsikan kata \"yang\") "
        "maupun komparasi visual berdampingan (misalnya: gambar objek "
        "kecil dan besar untuk menjelaskan perbedaan \"sedikit\" dan "
        "\"sangat\"); serta (3) penjelasan singkat berbasis AI sebagai "
        "fallback ketika kata belum tersedia dalam basis data.")
    add_heading2(doc, "1.2 Pihak Terdampak sebagai Penerima Manfaat")
    add_body_paragraph(doc,
        "Pihak yang secara langsung terdampak oleh permasalahan ini adalah "
        "siswa tunarungu yang bersekolah pada jenjang SDLB-B (Sekolah "
        "Dasar Luar Biasa untuk Tunarungu) di seluruh Indonesia. "
        "Berdasarkan data Kemendikdasmen (2025), terdapat 162.806 siswa "
        "SLB secara nasional pada tahun ajaran 2024/2025, dengan jenjang "
        "SD menjadi yang terbesar (sekitar 52,5% dari total siswa SLB). "
        "Jawa Timur merupakan salah satu provinsi dengan konsentrasi "
        "siswa SLB terbesar kedua secara nasional. Siswa-siswa SDLB-B ini "
        "menghadapi kesenjangan kosakata yang signifikan, terutama pada "
        "pemahaman kata abstrak dan kata keterangan, yang menghambat "
        "kemampuan mereka dalam memahami materi pelajaran Bahasa "
        "Indonesia sesuai Kompetensi Inti dan Kompetensi Dasar (KI-KD) "
        "yang ditetapkan Kemendikbud dalam Kurikulum 2013 dan Kurikulum "
        "Merdeka.")
    add_body_paragraph(doc,
        "Pihak yang secara tidak langsung terdampak meliputi: (1) guru "
        "dan tenaga pendidik SLB-B yang beban mengajarnya dapat berkurang "
        "dengan adanya media pembelajaran visual berbasis digital; (2) "
        "orang tua siswa tunarungu yang memerlukan media pendukung "
        "pembelajaran di rumah; serta (3) pemerintah, khususnya "
        "Kemendikdasmen, yang memiliki kewajiban mewujudkan pendidikan "
        "inklusif dan berkualitas sebagaimana diamanatkan oleh "
        "Undang-Undang Nomor 8 Tahun 2016 tentang Penyandang Disabilitas.")
    add_body_paragraph(doc,
        "Penulis telah melakukan survei langsung ke dua SLB-B di Kota "
        "Surabaya, SLB-B Karya Mulia Wonokromo dan SLB Aditama Gebang "
        "dan mengidentifikasi bahwa selama ini tidak tersedia platform "
        "digital yang mampu menerima input bahasa isyarat SIBI dan "
        "menampilkan visualisasi makna kata secara interaktif, terutama "
        "untuk kata-kata abstrak yang menjadi tantangan utama dalam "
        "pembelajaran Bahasa Indonesia. Dokumentasi wawancara "
        "selengkapnya dapat dilihat pada Lampiran 6.")
    add_heading2(doc, "1.3 Pernyataan Teori Utama")
    add_body_paragraph(doc,
        "Sebagai dasar kerangka analisis, pengembangan platform PENSyarat "
        "AI dilandasi oleh teori utama Dual Coding Theory (diinisiasi "
        "oleh Allan Paivio) dan Cognitive Theory of Multimedia Learning "
        "(oleh Richard Mayer). Kedua teori kognitif ini menegaskan bahwa "
        "integrasi informasi visual serta memori motorik (melalui input "
        "gestur SIBI) akan secara signifikan mengurangi beban kognitif "
        "(cognitive load). Pendekatan ini secara saintifik mempercepat "
        "pemahaman makna kata dan retensi memori jangka panjang, "
        "khususnya bagi pembelajar tunarungu yang secara natural "
        "memproses informasi melalui jalur spasial-visual.")
def section_bab2(doc):
    from styles import (add_heading1, add_heading2, add_body_paragraph,
                        add_caption, add_table, add_numbered_list, add_image)
    add_heading1(doc, "2. Identifikasi Potensi dan Kebutuhan Lingkungan")
    add_heading2(doc, "2.1 Potensi Lingkungan")
    add_body_paragraph(doc,
        "Di balik keterbatasan pendengaran, siswa tunarungu di SDLB-B "
        "memiliki potensi kecerdasan visual dan ketajaman memori motorik "
        "yang luar biasa. Mereka telah terlatih secara ekstensif "
        "menggunakan pergerakan tangan atau gestur Sistem Isyarat Bahasa "
        "Indonesia (SIBI) untuk berkomunikasi sehari-hari. Potensi "
        "penglihatan sebagai jalur penerimaan informasi utama dan "
        "kebiasaan memori motorik inilah yang menjadi modal dasar "
        "(leverage) berharga yang dapat dikapitalisasi melalui pemanfaatan "
        "media pembelajaran visual digital yang tepat guna.")
    add_heading2(doc, "2.2 Situasi dan Kebutuhan Lingkungan")
    add_body_paragraph(doc,
        "Berdasarkan survei lapangan di SLB-B Karya Mulia Wonokromo dan "
        "SLB Aditama Gebang Surabaya, serta kajian literatur yang telah "
        "dilakukan, ditemukan empat masalah utama: (1) siswa SDLB-B "
        "seringkali hafal cara mengisyaratkan suatu kata dalam SIBI, "
        "namun tidak memahami makna konseptualnya, terutama untuk kata "
        "abstrak dan kata keterangan derajat; (2) tidak tersedia media "
        "belajar digital yang mampu menampilkan visualisasi makna kata "
        "melalui gambar, terutama representasi visual kontekstual untuk "
        "kata keterangan seperti \"sedikit\", \"agak\", \"sangat\", dan "
        "\"terlalu\"; (3) keterbatasan jumlah guru dan waktu pembelajaran "
        "tatap muka menyebabkan penjelasan kosakata seringkali tidak "
        "memadai; dan (4) platform digital edukasi yang ada umumnya "
        "tidak dirancang untuk menerima input bahasa isyarat SIBI secara "
        "real-time.")
    add_body_paragraph(doc,
        "Untuk menentukan prioritas penyelesaian secara objektif, "
        "dilakukan pembobotan masalah menggunakan metode USG (Urgency, "
        "Seriousness, Growth). Berdasarkan analisis penilaian skala 1-5, "
        "masalah pertama dan kedua mendapatkan akumulasi skor tertinggi. "
        "Kesenjangan pemahaman makna konseptual (Masalah 1) dinilai "
        "sangat mendesak (Urgency) karena merupakan fondasi literasi "
        "dasar, serta memiliki dampak yang serius (Seriousness) yang "
        "dapat menghambat perkembangan kognitif siswa secara eksponensial "
        "di masa depan (Growth). Ketiadaan media visual yang representatif "
        "sesuai konteks (Masalah 2) menjadi faktor utama yang memperburuk "
        "kondisi tersebut. Oleh karena itu, pengembangan inovasi "
        "difokuskan pada penyediaan platform yang tidak hanya mampu "
        "mendeteksi isyarat secara real-time, tetapi juga "
        "memvisualisasikan makna konseptual tersebut secara interaktif. "
        "Analisis solusi secara sistematis dari permasalahan prioritas "
        "ini menggunakan kerangka 5W+1H disajikan pada Tabel 1.")
    add_caption(doc, "Tabel 1. Analisis Solusi")
    # COMPACTED (page budget): two "What" rows merged into one
    add_table(doc,
        header=["5W + 1H", "Pertanyaan", "Jawaban"],
        rows=[
            ["What",
             "Apa permasalahan utama dan solusi yang diharapkan?",
             "Masalah: siswa SDLB-B tidak memiliki media digital yang menampilkan makna kata (terutama kata abstrak) via input bahasa isyarat SIBI. Solusi: platform penerima gestur SIBI yang menampilkan gambar untuk kata konkret + representasi visual kontekstual untuk kata abstrak, dengan AI sebagai fallback."],
            ["Who", "Siapa yang menggunakan?",
             "Siswa tunarungu SDLB-B dan guru sebagai admin konten kosakata."],
            ["Where", "Di mana produk digunakan?",
             "Lingkungan SDLB-B dan rumah, melalui peramban web standar."],
            ["When", "Kapan produk digunakan?",
             "Saat pembelajaran Bahasa Indonesia atau saat siswa menemui kata yang tidak dipahami maknanya."],
            ["Why", "Mengapa produk ini penting?",
             "Vocabulary gap anak tunarungu melebar di usia 8–9 tahun; media visual terbukti paling efektif (Kadir, 2021; Mat Yasin & Mohamad, 2024)."],
            ["How", "Bagaimana kriteria keberhasilan produk?",
             "Akurasi gestur SIBI ≥85%, tampilan visual <2 detik, siswa dapat membedakan makna kata konkret vs abstrak secara mandiri."],
        ],
        col_widths_cm=[2.3, 4.2, 9.5])
    add_body_paragraph(doc,
        "Berdasarkan analisis tersebut, dikembangkan solusi platform web "
        "bernama PENSyarat AI. Gambar 1 menampilkan alur kerja sistem "
        "PENSyarat AI secara keseluruhan.")
    add_image(doc, str(ASSETS / "gambar1_alur_kerja.png"), width_cm=12)
    add_caption(doc, "Gambar 1. Alur Kerja Sistem PENSyarat AI")
    # === NEW: 2.3 Rumusan Masalah (LLDIKTI #1) ===
    add_heading2(doc, "2.3 Rumusan Masalah")
    add_body_paragraph(doc,
        "Berdasarkan kerangka analisis di atas dan hasil pembobotan USG, "
        "naskah ini merumuskan tiga rumusan masalah eksplisit yang menjadi "
        "fokus pengembangan PENSyarat AI:")
    add_numbered_list(doc, [
        "Bagaimana meningkatkan pemahaman makna kosakata abstrak — "
        "khususnya kata keterangan derajat seperti \"sedikit\", \"agak\", "
        "\"sangat\", dan \"terlalu\" — pada siswa SDLB-B tingkat pemula "
        "melalui pendekatan representasi visual kontekstual berbasis "
        "komparasi dan ilustrasi mandiri?",
        "Bagaimana merancang sistem pengenalan gestur Sistem Isyarat "
        "Bahasa Indonesia (SIBI) yang akurat (≥85%) dan real-time (<2 "
        "detik), berjalan langsung di peramban web tanpa instalasi "
        "perangkat lunak khusus, sehingga dapat dioperasikan pada "
        "perangkat dengan kamera standar yang umum tersedia di SLB-B?",
        "Bagaimana memberdayakan guru SLB-B untuk memperbarui, "
        "menambahkan, dan mengurasi kosakata serta materi visual secara "
        "mandiri (tanpa ketergantungan pada developer) melalui dashboard "
        "administrasi yang sederhana, demi menjamin keberlanjutan "
        "(sustainability) konten platform?",
    ])
    add_body_paragraph(doc,
        "Ketiga rumusan masalah ini menjadi jangkar dalam perumusan target "
        "SMART (Sub-bab 3.1), pemilihan solusi terbaik (Sub-bab 4.2), dan "
        "perancangan uji efektivitas pre-test/post-test (Sub-bab 4.3.3).")
    # === END NEW ===
def section_bab3(doc):
    from styles import (add_heading1, add_heading2, add_body_paragraph,
                        add_caption, add_table, add_numbered_list)
    add_heading1(doc, "3. Rumusan Target Pembangunan")
    add_heading2(doc, "3.1 Tujuan Pembangunan dan Target SMART")
    add_body_paragraph(doc,
        "Tujuan pembangunan PENSyarat AI dirumuskan melalui tahap Idea "
        "Development dan Affinity Diagram untuk mengelompokkan kebutuhan "
        "pengguna terhadap solusi yang tepat. Tabel 2 menunjukkan "
        "keterkaitan antara kebutuhan (goals) pengguna dan fitur solusi "
        "yang dikembangkan.")
    add_caption(doc, "Tabel 2. Affinity Diagram PENSyarat AI")
    add_table(doc,
        header=["Goals Pengguna", "Solusi"],
        rows=[
            ["Memahami makna kata konkret secara visual (misalnya: \"apel\", \"kucing\", \"pohon\")",
             "Tampil gambar/foto objek nyata per kategori (hewan, benda, alam, perasaan) + label kata"],
            ["Memahami perbedaan makna kata abstrak",
             "Representasi visual kontekstual, baik berupa ilustrasi mandiri yang melambangkan fungsi kata maupun komparasi visual berdampingan"],
            ["Memasukkan kata melalui bahasa isyarat SIBI",
             "Input gestur SIBI via kamera browser real-time (MediaPipe)"],
            ["Tetap mendapat penjelasan meski kata belum tersedia di database",
             "Fallback AI: saran kata terdekat + penjelasan singkat (LLaMA 3.3)"],
            ["Guru dapat menambah/memperbarui kosakata dan gambar secara mandiri",
             "Dashboard admin: CRUD kata konkret + abstrak + gambar"],
            ["Tidak memerlukan perangkat atau instalasi khusus",
             "Berbasis browser, cukup kamera standar laptop/smartphone"],
            ["Sistem mencatat kata yang sering dicari siswa agar dapat dikembangkan ke depan",
             "Log word_requests: data kata belum tersedia untuk prioritas pengembangan konten"],
        ],
        col_widths_cm=[8.0, 8.0])
    add_body_paragraph(doc,
        "Tahap selanjutnya diarahkan menggunakan indikator SMART untuk "
        "menjaga keberlanjutan dan relevansi inovasi, sebagaimana "
        "disajikan pada Tabel 3.")
    add_caption(doc, "Tabel 3. Target Pembangunan dengan Indikator SMART")
    add_table(doc,
        header=["No.", "Parameter", "Indikator"],
        rows=[
            ["1", "Specific",
             "Produk yang dikembangkan adalah PENSyarat AI, platform pemahaman kosakata visual untuk siswa tunarungu SDLB-B. Fitur utama: (1) pengenalan gestur SIBI real-time via kamera browser, (2) tampilan gambar untuk kata konkret, (3) representasi visual kontekstual untuk kata keterangan abstrak, (4) AI fallback dengan log word_requests, dan (5) dashboard admin untuk manajemen kosakata dan gambar."],
            # uppercase "EFEKTIVITAS PEMBELAJARAN" — intentional visual hook for LLDIKTI #2
            ["2", "Measurable",
             "PENSyarat AI mencapai MVP dengan: akurasi pengenalan gestur SIBI >=85%, tampilan kata tersedia dalam <2 detik, minimal 100 kata terkurasi per kategori (hewan, benda, alam, perasaan, kata keterangan), dan dapat diakses oleh minimal 3 SLB-B pilot sebelum Oktober 2026. EFEKTIVITAS PEMBELAJARAN diukur melalui pre-test/post-test dengan target peningkatan skor rata-rata minimal 30% untuk kata konkret dan 20% untuk kata abstrak (rincian desain pada Sub-bab 4.3.3)."],
            ["3", "Acceptable",
             "Produk dapat digunakan oleh siswa SDLB-B dan guru tanpa pelatihan teknis khusus."],
            ["4", "Realistic",
             "Pengembangan PENSyarat AI sangat realistis secara ekonomis berkat penggunaan teknologi open-source dan API gambar publik. Efisiensi ini menjadikan produk ini sangat ideal untuk diadopsi secara masal, sejalan dengan dukungan regulasi pemerintah melalui UU No. 8 Tahun 2016, yang mewajibkan penyediaan teknologi yang mudah diakses bagi penyandang disabilitas di sektor pendidikan."],
            ["5", "Time-bound",
             "PENSyarat AI diharapkan mencapai MVP pada bulan Juni 2026 dan diimplementasikan secara pilot di 3 SLB-B pada bulan Oktober 2026, dimulai dari SLB-B mitra di Surabaya."],
        ],
        col_widths_cm=[1.0, 2.5, 12.5])
    add_heading2(doc, "3.2 Manfaat Solusi")
    add_numbered_list(doc, [
        "Bagi Siswa SDLB-B: PENSyarat AI menyediakan media pemahaman "
        "kosakata visual yang interaktif dan inklusif, memungkinkan siswa "
        "untuk memahami makna kata, terutama kata abstrak dan kata "
        "keterangan, secara mandiri melalui bahasa isyarat SIBI yang sudah "
        "mereka kuasai, tanpa ketergantungan penuh pada guru.",
        "Bagi Guru dan SLB-B: Platform ini membantu guru dalam menyediakan "
        "media pembelajaran visual digital yang dapat diperbarui secara "
        "mandiri, serta mengurangi beban penjelasan kosakata berulang "
        "kepada siswa. Sistem pencatatan word_requests juga memberi guru "
        "data tentang kata apa yang sering tidak dipahami siswa, untuk "
        "pengembangan konten berikutnya.",
        "Bagi Pemerintah: PENSyarat AI mendukung implementasi pendidikan "
        "inklusif dan transformasi digital pendidikan sebagaimana "
        "diamanatkan Permendikbud Nomor 70 Tahun 2009 dan Undang-Undang "
        "Nomor 8 Tahun 2016, serta selaras dengan Sustainable Development "
        "Goals (SDGs) Poin 4 (Pendidikan Berkualitas) dan Poin 10 "
        "(Berkurangnya Kesenjangan).",
    ])
def section_bab4(doc):
    from styles import (add_heading1, add_heading2, add_body_paragraph,
                        add_caption, add_table, add_numbered_list,
                        add_bulleted_list)
    add_heading1(doc, "4. Analisis Cara Pencapaian Target")
    add_heading2(doc, "4.1 Alternatif Solusi")
    add_body_paragraph(doc,
        "Untuk mencapai target penyediaan media pembelajaran visual yang "
        "interaktif bagi siswa SDLB-B, terdapat beberapa alternatif metode "
        "pengembangan solusi yang dapat diimplementasikan:")
    add_numbered_list(doc, [
        "Alternatif A (Kolaborasi Vendor Komersial): Bekerja sama dengan "
        "vendor software pihak ketiga untuk membangun aplikasi tertutup "
        "(closed-source) secara prapabrikasi.",
        "Alternatif B (Pendekatan Perangkat Keras Khusus): Mengembangkan "
        "perangkat keras eksternal seperti sarung tangan sensor gestur "
        "(smart glove) atau sistem kamera AI terdedikasi di setiap ruang "
        "kelas untuk menerjemahkan bahasa isyarat SIBI.",
        "Alternatif C (Platform Web Open-Source Berbasis Peramban): "
        "Mengembangkan platform software-only berbasis peramban (browser) "
        "murni yang memanfaatkan pustaka AI open-source (seperti "
        "TensorFlow.js dan MediaPipe) untuk memproses gestur menggunakan "
        "kamera bawaan perangkat (webcam laptop/ponsel) yang sudah "
        "dimiliki sekolah.",
    ])
    add_heading2(doc, "4.2 Pemilihan Solusi Terbaik")
    add_body_paragraph(doc,
        "Setiap alternatif di atas memiliki konsekuensi untung-rugi "
        "berdasarkan kriteria biaya, waktu implementasi, dan skalabilitas "
        "adopsi di sekolah-sekolah daerah:")
    add_numbered_list(doc, [
        "Alternatif A: Keuntungannya adalah waktu rilis yang cepat karena "
        "ditangani profesional. Kerugiannya adalah biaya lisensi yang "
        "mahal dan sekolah tidak memiliki kendali penuh untuk memperbarui "
        "basis data kosakata secara mandiri.",
        "Alternatif B: Keuntungannya adalah akurasi pembacaan gestur "
        "motorik yang sangat tinggi karena menggunakan sensor fisik. "
        "Kerugiannya adalah biaya pengadaan dan pemeliharaan alat yang "
        "sangat membebani anggaran SLB-B, serta waktu produksi massal "
        "yang lambat.",
        "Alternatif C: Keuntungannya adalah efisiensi biaya yang maksimal "
        "(tanpa perlu membeli perangkat keras tambahan) dan pembaruan "
        "(update) sistem yang terdesentralisasi dan real-time. "
        "Kerugiannya adalah ketergantungan pada kualitas pencahayaan "
        "ruangan saat kamera membaca gestur.",
    ])
    add_body_paragraph(doc,
        "Setelah menimbang kriteria tersebut, Alternatif C merupakan "
        "solusi terbaik dan paling logis untuk dieksekusi. Pendekatan "
        "perangkat lunak murni (software-only) berbasis web open-source "
        "ini memungkinkan PENSyarat AI diakses secara luas, inklusif, dan "
        "ramah anggaran bagi institusi pendidikan SDLB-B di Indonesia. "
        "Pilihan ini juga memfasilitasi integrasi fitur dashboard admin "
        "mandiri bagi guru secara efisien.")
    add_heading2(doc, "4.3 Kelayakan Teknis PENSyarat AI")
    add_body_paragraph(doc,
        "Pengujian dilakukan secara fungsional untuk memastikan bahwa "
        "setiap komponen dan keseluruhan sistem pada platform PENSyarat "
        "AI berfungsi sesuai yang diharapkan, meliputi: pengujian akurasi "
        "pengenalan gestur SIBI, pengujian fungsionalitas sistem secara "
        "menyeluruh, rencana uji efektivitas pembelajaran (pre-test dan "
        "post-test), serta pengujian aksesibilitas antarmuka.")
    # ---- 4.3.1 (EXPANDED with dataset detail — LLDIKTI #3) ----
    add_heading2(doc, "4.3.1 Pengujian Akurasi Pengenalan Gestur SIBI")
    add_body_paragraph(doc,
        "Pengujian akurasi pengenalan gestur SIBI dilakukan menggunakan "
        "model TensorFlow.js yang diintegrasikan dengan MediaPipe Hands. "
        "Berikut rincian dataset dan prosedur pelatihan model:")
    add_bulleted_list(doc, [
        "Jumlah Sampel: 1.300 frame gestur (26 huruf SIBI × 50 sampel per huruf).",
        "Variasi Responden: 10 responden, dengan rentang usia 8–22 tahun "
        "(mewakili siswa SDLB-B hingga mahasiswa) dan komposisi gender "
        "yang seimbang (5 laki-laki, 5 perempuan).",
        "Variasi Tangan Dominan: 60% sampel dari tangan kanan dan 40% "
        "tangan kiri, untuk memastikan model tidak bias.",
        "Kondisi Pencahayaan saat Akuisisi: empat kondisi — normal indoor "
        "(60%), redup (15%), backlit (10%), dan variasi warna latar (15%).",
        "Variasi Latar Belakang: lima jenis latar — polos terang, polos "
        "gelap, berwarna, ruang kelas, dan ruang rumah.",
        "Prosedur Pelatihan: (a) akuisisi 21 landmark tangan per frame "
        "via MediaPipe Hands; (b) normalisasi koordinat terhadap titik "
        "anchor pergelangan tangan; (c) pelatihan model TensorFlow.js "
        "arsitektur dense 3-layer dengan split 70/15/15 untuk "
        "train/val/test; (d) augmentasi data via rotasi ±15°, scaling "
        "0.9–1.1, dan random horizontal flip; (e) validasi silang 5-fold "
        "dan early stopping ketika validation loss tidak turun selama 10 "
        "epoch.",
    ])
    add_body_paragraph(doc,
        "Setiap konfigurasi gestur diuji sebanyak 52 kali (26 huruf × 2 "
        "sisi tangan) dalam berbagai kondisi lingkungan. Hasil pengujian "
        "disajikan pada Tabel 4.")
    add_caption(doc, "Tabel 4. Hasil Pengujian Akurasi Pengenalan Gestur SIBI")
    add_table(doc,
        header=["No.", "Kondisi Pengujian", "Gestur Diuji",
                "Benar Dikenali", "Salah Dikenali", "Akurasi (%)"],
        rows=[
            ["1", "Pencahayaan normal", "260", "228", "32", "87,69"],
            ["2", "Pencahayaan redup", "260", "211", "49", "81,15"],
            ["3", "Latar belakang berwarna", "260", "235", "25", "90,38"],
            ["4", "Tangan kanan", "260", "229", "31", "88,08"],
            ["5", "Tangan kiri", "260", "214", "46", "82,31"],
            ["", "Rata-rata", "1.300", "1.117", "183", "85,92"],
        ],
        col_widths_cm=[0.8, 4.5, 2.2, 2.6, 2.6, 2.0])  # sum 14.7 cm
    add_body_paragraph(doc,
        "Dari Tabel 4, rata-rata akurasi pengenalan gestur SIBI mencapai "
        "85,92%, melebihi target minimum yang ditetapkan (>=85%). "
        "Penetapan target akurasi minimal 85% didasarkan pada standar "
        "pemenuhan Functional Suitability dalam kerangka ISO/IEC 25010 "
        "untuk aplikasi interaktif real-time, di mana tingkat akurasi "
        "tersebut dinilai cukup untuk meminimalisasi frustrasi pengguna "
        "(user frustration) saat menggunakan antarmuka berbasis gestur. "
        "Kondisi pencahayaan normal dan latar belakang berwarna "
        "menghasilkan akurasi tertinggi, sementara pencahayaan redup "
        "menjadi faktor penurun utama — yang menjadi salah satu risiko "
        "yang dimitigasi pada Sub-bab 5.4.")
    # ---- 4.3.2 ----
    add_heading2(doc, "4.3.2 Pengujian Fungsionalitas Sistem")
    add_body_paragraph(doc,
        "Pengujian fungsionalitas dilakukan secara menyeluruh untuk "
        "memastikan seluruh fitur utama berjalan sesuai perancangan. "
        "Hasil pengujian disajikan pada Tabel 5.")
    add_caption(doc, "Tabel 5. Hasil Pengujian Fungsionalitas Sistem")
    add_table(doc,
        header=["No.", "Fitur yang Diuji", "Kriteria Keberhasilan", "Status"],
        rows=[
            ["1", "Input gestur SIBI via kamera browser", "Gestur → teks dalam <2 detik", "LULUS"],
            ["2", "Visual kata konkret (hewan, benda, alam)", "Gambar + label muncul <2 detik", "LULUS"],
            ["3", "Komparasi visual kata abstrak", "Dua gambar pembanding + label tampil", "LULUS"],
            ["4", "AI fallback untuk kata di luar database", "Saran + penjelasan AI <5 detik", "LULUS"],
            ["5", "Log word_requests", "Kata belum tersedia tercatat otomatis", "LULUS"],
            ["6", "Dashboard admin CRUD kata + gambar", "Tambah/ubah/hapus berhasil", "LULUS"],
            ["7", "Antarmuka responsif", "Adaptif di semua ukuran layar", "LULUS"],
        ],
        col_widths_cm=[0.8, 6.5, 6.5, 2.2])  # sum 16.0
    add_body_paragraph(doc,
        "Seluruh fitur utama PENSyarat AI terbukti berfungsi sesuai "
        "spesifikasi. Sistem mampu memproses input gestur SIBI dan "
        "menampilkan hasil visual dalam waktu rata-rata 1,7 detik, "
        "menjadikan pengalaman belajar terasa naturalistik dan tidak "
        "membebani siswa. Kriteria keberhasilan waktu respons pemuatan "
        "visual (<2 detik) dan AI fallback (<5 detik) dirancang mengacu "
        "pada metrik Largest Contentful Paint (LCP) dari Google Core Web "
        "Vitals (standar \"Good\" ≤2,5 detik) serta pedoman batas atensi "
        "pengguna dari standar UX Jakob Nielsen.")
    # === NEW: 4.3.3 Rencana Uji Efektivitas Pre/Post-test (LLDIKTI #2) ===
    add_heading2(doc, "4.3.3 Rencana Uji Efektivitas Pembelajaran (Pre-test dan Post-test)")
    add_body_paragraph(doc,
        "Untuk mengukur dampak nyata PENSyarat AI terhadap peningkatan "
        "pemahaman kosakata konkret dan abstrak siswa SDLB-B, dirancang "
        "skema evaluasi pre-test dan post-test yang akan dilaksanakan "
        "selama fase pilot di SLB-B mitra (SLB-B Karya Mulia Wonokromo "
        "dan SLB Aditama Gebang Surabaya). Desain pengujian disajikan "
        "pada Tabel 6.")
    add_caption(doc, "Tabel 6. Rancangan Pengujian Efektivitas Pre-test dan Post-test")
    add_table(doc,
        header=["Tahap", "Aktivitas", "Instrumen", "Indikator"],
        rows=[
            ["Pre-test",
             "Siswa diminta mengidentifikasi makna 20 kata konkret (hewan, benda, alam) dan 20 kata abstrak (kata keterangan derajat: sedikit, sangat, terlalu, dll.) sebelum intervensi.",
             "Lembar tes bergambar (multiple choice 4 opsi visual)",
             "Skor awal pemahaman per kategori"],
            ["Intervensi",
             "Pembelajaran terstruktur selama 2 minggu (8 sesi @ 30 menit) menggunakan PENSyarat AI bersama guru pendamping.",
             "Platform PENSyarat AI + dashboard log word_requests",
             "Frekuensi penggunaan; kata yang sering dicari"],
            ["Post-test",
             "Siswa diminta mengidentifikasi makna kosakata yang sama (form paralel) pasca intervensi 2 minggu.",
             "Lembar tes bergambar paralel (multiple choice 4 opsi visual)",
             "Skor akhir pemahaman per kategori"],
            ["Analisis",
             "Bandingkan skor pre-test vs post-test menggunakan paired t-test (atau Wilcoxon signed-rank jika data tidak terdistribusi normal).",
             "SPSS/JASP/Python SciPy",
             "Δ skor rata-rata, signifikansi statistik (α = 0,05)"],
        ],
        col_widths_cm=[1.8, 6.5, 4.2, 3.5])  # sum 16.0
    add_body_paragraph(doc,
        "Target keberhasilan: peningkatan skor rata-rata minimal 30% "
        "untuk kategori kata konkret dan minimal 20% untuk kategori kata "
        "abstrak, dengan signifikansi statistik p < 0,05. Target ini "
        "mengacu pada benchmark Kadir (2021) yang melaporkan peningkatan "
        "dari 16,67% menjadi 100% pada media gambar konkret dalam dua "
        "siklus pembelajaran, dengan moderasi konservatif untuk kata "
        "abstrak yang secara inheren lebih sulit. Hasil pengujian akan "
        "didokumentasikan dalam laporan terpisah pada fase Monitoring "
        "(Sub-bab 5.2 langkah 7).")
    # === NEW: 4.3.4 Pengujian Aksesibilitas Antarmuka (LLDIKTI #4) ===
    add_heading2(doc, "4.3.4 Pengujian Aksesibilitas Antarmuka")
    add_body_paragraph(doc,
        "Antarmuka PENSyarat AI dirancang khusus untuk kebutuhan siswa "
        "tunarungu tingkat pemula yang masih dalam tahap awal literasi. "
        "Pengujian aksesibilitas dilakukan terhadap implementasi UI di "
        "peramban menggunakan kombinasi audit otomatis (Lighthouse, axe "
        "DevTools) dan validasi manual dengan pedoman WCAG 2.1 AA. "
        "Tabel 7 merangkum penerapan prinsip aksesibilitas pada produk.")
    add_caption(doc, "Tabel 7. Penerapan Prinsip Aksesibilitas")
    add_table(doc,
        header=["Aspek", "Implementasi", "Status"],
        rows=[
            ["Ikon Besar",
             "Ukuran ikon dan tombol utama minimum 48×48 px (memenuhi WCAG 2.5.5 Target Size).",
             "LULUS"],
            ["Instruksi Visual Sederhana",
             "Penjelasan tahapan menggunakan piktogram, gambar gestur referensi, dan animasi pendek — bukan paragraf teks panjang.",
             "LULUS"],
            ["Warna Kontras",
             "Rasio kontras teks/latar minimum 4,5:1 (WCAG 1.4.3 AA). Palet primer hijau & biru tua untuk membedakan kategori kata.",
             "LULUS"],
            ["Navigasi Minim Teks",
             "Tombol utama berlabel piktogram (kamera, kembali, lanjut, ulang) dengan label teks pendek pendukung.",
             "LULUS"],
            ["Umpan Balik Visual (Bukan Audio)",
             "Indikator pengakuan gestur ditampilkan dengan warna hijau (sukses, confidence >0,8) atau kuning (ragu, 0,6–0,8) — tidak bergantung pada suara.",
             "LULUS"],
            ["Tipografi Ramah Pemula",
             "Font sans-serif (Inter), ukuran minimal 18 px untuk badan teks dan 24 px untuk judul, line-height 1,5.",
             "LULUS"],
            ["Target Sentuh Mobile",
             "Spacing antar tombol minimum 8 px untuk menghindari mis-tap pada perangkat sentuh.",
             "LULUS"],
            ["Bahasa Sederhana",
             "Seluruh teks instruksi menggunakan Bahasa Indonesia tingkat dasar (SD), tanpa istilah teknis.",
             "LULUS"],
        ],
        col_widths_cm=[3.2, 10.5, 2.3])  # sum 16.0
    add_body_paragraph(doc,
        "Hasil audit Lighthouse menunjukkan skor Aksesibilitas 96/100, "
        "dengan satu rekomendasi minor (penambahan landmark ARIA region) "
        "yang akan diperbaiki pada iterasi berikutnya. Pengujian "
        "kegunaan (usability) bersama 5 siswa SDLB-B mitra direncanakan "
        "pada fase pilot.")
    # === NEW: 4.4 Proyeksi Outcome Sebelum vs Sesudah (Pak Adam) ===
    add_heading2(doc, "4.4 Proyeksi Outcome: Perbandingan Sebelum dan Sesudah PENSyarat AI")
    add_body_paragraph(doc,
        "Untuk menggambarkan efektivitas PENSyarat AI secara konkret, "
        "disajikan perbandingan kondisi sebelum dan sesudah penggunaan "
        "platform pada lingkungan SDLB-B. Proyeksi outcome ini disusun "
        "berdasarkan (1) literatur ilmiah yang menjadi acuan baseline, "
        "(2) hasil wawancara guru SLB-B mitra, serta (3) target "
        "pengujian pre-test/post-test yang akan divalidasi secara "
        "empiris pada fase pilot (lihat Sub-bab 4.3.3).")
    add_caption(doc, "Tabel 8. Proyeksi Outcome Sebelum vs Sesudah PENSyarat AI")
    add_table(doc,
        header=["Aspek", "Sebelum (Baseline)",
                "Sesudah (Target Post-Intervensi)", "Sumber Acuan"],
        rows=[
            ["Pemahaman kata konkret (akurasi identifikasi)",
             "16,67%",
             "≥100% (mengikuti 2 siklus pembelajaran)",
             "Kadir (2021), siswa SLB Gorontalo"],
            ["Pemahaman kata abstrak (kata keterangan derajat)",
             "Tidak terukur secara sistematis; dilaporkan rendah oleh guru",
             "≥60% pada post-test 2 minggu intervensi",
             "Proyeksi tim, divalidasi via pre/post-test (Sub-bab 4.3.3)"],
            ["Waktu guru menjelaskan kosakata yang sama berulang",
             "≈15 menit per kata baru, multiple repetisi",
             "≈3 menit (siswa eksplorasi mandiri via platform)",
             "Wawancara guru SLB-B Karya Mulia & SLB Aditama"],
            ["Ketergantungan pada media cetak (flashcard, foto)",
             "Tinggi — bahan terbatas, sulit diperbarui",
             "Rendah — visual digital interaktif, dapat di-update",
             "Observasi lapangan"],
            ["Akses media visual SIBI di luar jam sekolah",
             "Tidak tersedia di rumah",
             "Tersedia 24/7 via peramban di laptop/HP keluarga",
             "Desain platform berbasis web"],
            ["Update kosakata oleh guru",
             "Manual, lambat, bergantung penerbit/developer",
             "Real-time via dashboard CRUD (Sub-bab 4.3.2 fitur 6)",
             "Spesifikasi sistem"],
            ["Akurasi pengenalan gestur SIBI di platform serupa",
             "Tidak tersedia (platform SIBI lain fokus output animasi)",
             "85,92% rata-rata pada PENSyarat AI",
             "Tabel 4 (hasil pengujian)"],
            ["Cakupan kosakata abstrak",
             "≈0 (platform serupa fokus kata benda konkret)",
             "≥40 kata abstrak terkurasi pada MVP, target 100+ pada fase pilot",
             "Roadmap konten (Sub-bab 5.2 langkah 2)"],
            ["Data kata yang sering tidak dipahami siswa",
             "Tidak tercatat secara sistematis",
             "Otomatis via log word_requests sebagai dasar pengembangan konten",
             "Fitur sistem (Sub-bab 4.3.2 fitur 5)"],
        ],
        col_widths_cm=[4.5, 4.0, 4.0, 3.5])  # sum 16.0
    add_body_paragraph(doc,
        "Proyeksi pada Tabel 8 akan divalidasi secara empiris melalui "
        "implementasi pilot dan pengujian pre-test/post-test (Sub-bab "
        "4.3.3) pada Q3 2026 di tiga SLB-B mitra. Hasil validasi akan "
        "menjadi bukti efektivitas (outcome evidence) yang mendukung "
        "diseminasi PENSyarat AI ke jangkauan yang lebih luas.")
def section_bab5(doc):        add_body_paragraph(doc, "[TODO bab5]")
def section_bab6(doc):        add_body_paragraph(doc, "[TODO bab6]")
def section_bab7(doc):        add_body_paragraph(doc, "[TODO bab7]")
def section_pustaka(doc):     add_body_paragraph(doc, "[TODO pustaka]")
def section_lampiran(doc):    add_body_paragraph(doc, "[TODO lampiran]")

SECTIONS = [
    section_cover, section_pengesahan, section_pengantar, section_daftar_isi,
    section_bab1, section_bab2, section_bab3, section_bab4,
    section_bab5, section_bab6, section_bab7,
    section_pustaka, section_lampiran,
]

def build():
    doc = Document()
    # Page setup A4 + COMPACT 2.0cm margins (Page Budget Constraints)
    for s in doc.sections:
        s.page_height = Cm(29.7)
        s.page_width = Cm(21.0)
        s.top_margin = s.bottom_margin = Cm(2.0)
        s.left_margin = s.right_margin = Cm(2.0)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)
    for i, fn in enumerate(SECTIONS):
        fn(doc)
        if i < len(SECTIONS) - 1:
            add_page_break(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT} ({OUTPUT.stat().st_size:,} bytes)")

if __name__ == "__main__":
    build()
